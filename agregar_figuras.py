import docx
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

LEYENDAS = [
    # Módulo 3.6.1
    "Figura 3.1: Interfaz web de inicio de sesión de usuario.",
    "Figura 3.2: Vista de ingreso de código OTP de 6 dígitos.",
    "Figura 3.3: Validación de código expirado o incorrecto.",
    "Figura 3.4: Implementación del método verificarCodigo() con hash_equals().",
    "Figura 3.5: Implementación del método login() y regeneración de sesión.",
    "Figura 3.6: Configuración del limitador de tasa auth-throttle.",
    # Módulo 3.6.2
    "Figura 3.7: Formulario para registro y dimensiones físicas de nuevo recinto.",
    "Figura 3.8: Panel de administración de recintos y aforo.",
    "Figura 3.9: Validación de aforo máximo excedido.",
    "Figura 3.10: Configuración de zonas tarifarias por rangos de filas.",
    "Figura 3.11: Panel de eventos activos y gestión de tarifas.",
    "Figura 3.12: Formulario de alta de nuevo evento.",
    "Figura 3.13: Interfaz de edición y carga de banner promocional.",
    "Figura 3.14: Asignación de precios base por zona del evento.",
    "Figura 3.15: Algoritmo de generación de matriz con guardia lockForUpdate().",
    "Figura 3.16: Lógica de creación de zonas dentro de transacción atómica.",
    # Módulo 3.6.3 y 3.6.4
    "Figura 3.17: Landing page del evento con detalle de zonas y precios.",
    "Figura 3.18: Mapa interactivo en Grid Cartesiano con temporizador de reserva activo.",
    "Figura 3.19: Selección interactiva de asientos en tiempo real.",
    "Figura 3.20: Pantalla de confirmación de compra con comprobantes y códigos QR.",
    "Figura 3.21: Retorno exitoso desde la pasarela Mercado Pago Checkout Pro.",
    "Figura 3.22: Panel del cliente \"Mis Boletos\" con listado de compras.",
    "Figura 3.23: Correo transaccional recibido en Mailpit con código QR inyectado.",
    "Figura 3.24: Template del Grid Cartesiano en Vue 3 (EventoCheckout.vue).",
    "Figura 3.25: Renderizado dinámico de escenario circular SVG (Arena 360°).",
    "Figura 3.26: Componente reactivo de resumen de orden y cálculo de precios.",
    "Figura 3.27: Despliegue de métodos de pago tras completar la reserva.",
    "Figura 3.28: Interfaz del Punto de Venta (POS) en Taquilla.",
    "Figura 3.29: Selección de asientos y cobro en efectivo dentro de la taquilla POS.",
    "Figura 3.30: Impresión del recibo térmico tras confirmar la venta en POS.",
    "Figura 3.31: Método procesarCompra() con bloqueo pesimista en CompraService.php.",
    "Figura 3.32: Método comprarEnTaquilla() dentro de CompraService.php.",
    "Figura 3.33: Método reservarAsientos() con expiración a 5 minutos en CompraService.php.",
    "Figura 3.34: Método webhookMercadoPago() con validación HMAC en CompraController.php.",
    "Figura 3.35: Método generarQrDataUri() en ConfirmacionCompraMail.php.",
    "Figura 3.36: Método urlConfirmacion() en ConfirmacionCompraMail.php.",
    # Módulo 3.6.5
    "Figura 3.37: Estado \"Up\" de los 5 contenedores en Docker Compose.",
    "Figura 3.38: Verificación de extensiones PHP activas en el contenedor app.",
    "Figura 3.39: Configuración del servicio MySQL en docker-compose.yml.",
    "Figura 3.40: Definición de volúmenes compartidos en modo solo lectura (:ro).",
    "Figura 3.41: Construcción multietapa (multi-stage build) en Dockerfile.",
]

DRAWING_TAG = qn('w:drawing')


def make_caption_paragraph(doc, texto):
    p_pie = doc.add_paragraph()
    p_pie.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p_pie.add_run(texto)
    run.font.name = 'Calibri'
    run.font.size = Pt(10)
    run.font.italic = True
    run.font.bold = True
    run.font.color.rgb = RGBColor(80, 80, 80)
    return p_pie


doc = docx.Document("3.6.docx")
fig_index = 0

# Usamos una copia estática de los párrafos originales: doc.add_paragraph()
# añade nuevos párrafos al final del cuerpo, por lo que iterar sobre
# doc.paragraphs en vivo sería inseguro.
paragraphs_originales = list(doc.paragraphs)

for paragraph in paragraphs_originales:
    num_imagenes = len(paragraph._element.findall('.//' + DRAWING_TAG))
    if num_imagenes == 0:
        continue

    insert_after = paragraph._element
    for _ in range(num_imagenes):
        if fig_index >= len(LEYENDAS):
            break
        p_pie = make_caption_paragraph(doc, LEYENDAS[fig_index])
        insert_after.addnext(p_pie._element)
        insert_after = p_pie._element
        fig_index += 1

doc.save("3.6_Con_Figuras.docx")
print(f"Proceso completado exitosamente. Se insertaron {fig_index} pies de figura "
      f"sin alterar el texto ni las imágenes.")
